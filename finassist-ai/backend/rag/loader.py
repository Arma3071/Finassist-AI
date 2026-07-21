"""Document loading and cleaning for the RAG ingestion pipeline.

Supports PDF, DOCX, TXT, and Markdown files. Extracted text is cleaned
(duplicate whitespace collapsed) and headings are heuristically detected
so downstream chunking/metadata can make use of document structure.
"""

import re
from dataclasses import dataclass, field
from pathlib import Path

from backend.utils.logging_config import get_logger

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}

_HEADING_PATTERN = re.compile(r"^(#{1,6}\s+.+|[A-Z][A-Za-z0-9 ,'&-]{3,80})$")


@dataclass
class LoadedDocument:
    """Container for extracted, cleaned document text and metadata."""

    filename: str
    text: str
    headings: list[str] = field(default_factory=list)
    extension: str = ""


class UnsupportedFileTypeError(ValueError):
    """Raised when a file extension is not supported for ingestion."""


class DocumentLoader:
    """Loads and cleans raw text from PDF, DOCX, TXT, and Markdown files."""

    def load(self, file_path: str | Path) -> LoadedDocument:
        """Load a single file from disk into a :class:`LoadedDocument`.

        Args:
            file_path: Path to the source file.

        Returns:
            A LoadedDocument with cleaned text and detected headings.

        Raises:
            UnsupportedFileTypeError: If the extension is not supported.
            FileNotFoundError: If the path does not exist.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(f"Unsupported file type: {ext}")

        logger.info("Loading document %s", path.name)
        raw_text = self._extract(path, ext)
        cleaned = self._clean_text(raw_text)
        headings = self._detect_headings(cleaned)

        return LoadedDocument(filename=path.name, text=cleaned, headings=headings, extension=ext)

    def _extract(self, path: Path, ext: str) -> str:
        if ext == ".pdf":
            return self._extract_pdf(path)
        if ext == ".docx":
            return self._extract_docx(path)
        # .txt / .md / .markdown
        return path.read_text(encoding="utf-8", errors="ignore")

    @staticmethod
    def _extract_pdf(path: Path) -> str:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)

    @staticmethod
    def _extract_docx(path: Path) -> str:
        import docx

        document = docx.Document(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    @staticmethod
    def _clean_text(text: str) -> str:
        """Collapse duplicate whitespace while preserving paragraph breaks."""
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _detect_headings(text: str) -> list[str]:
        """Heuristically detect headings: markdown '#' lines or short Title-Case lines."""
        headings = []
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped or len(stripped) > 90:
                continue
            if _HEADING_PATTERN.match(stripped):
                headings.append(stripped.lstrip("#").strip())
        return headings
